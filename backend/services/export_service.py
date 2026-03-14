"""
Export service for generating deliverable files.

Handles TXT, Markdown, and DOCX export generation from database posts.
"""

import json
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from sqlalchemy.orm import Session

from backend.models.client import Client
from backend.models.post import Post
from backend.models.project import Project
from backend.utils.logger import logger


async def generate_export_file(
    posts: List[Post],
    client: Client,
    project: Project,
    format: str,
    relative_path: str,
    include_audit_log: bool = False,
    include_research: bool = False,
    db: Optional[Session] = None,
) -> Tuple[Path, int]:
    """
    Generate export file in specified format.

    Args:
        posts: List of Post database models
        client: Client database model
        project: Project database model
        format: Export format ('txt', 'md', or 'docx')
        relative_path: Relative path for the output file (from data/outputs/)
        include_audit_log: Whether to include audit log in export
        include_research: Whether to include research results appendix
        db: Database session (required if include_audit_log or include_research is True)

    Returns:
        Tuple of (absolute file path, file size in bytes)
    """
    # Ensure output directory exists
    output_dir = Path("data/outputs")
    full_path = output_dir / relative_path
    full_path.parent.mkdir(parents=True, exist_ok=True)

    if format == "docx":
        return await _generate_docx(
            posts, client, project, full_path, include_audit_log, include_research, db
        )
    elif format == "md" or format == "markdown":
        return await _generate_markdown(
            posts, client, project, full_path, include_audit_log, include_research, db
        )
    else:
        return await _generate_txt(
            posts, client, project, full_path, include_audit_log, include_research, db
        )


async def _generate_txt(
    posts: List[Post],
    client: Client,
    project: Project,
    output_path: Path,
    include_audit_log: bool,
    include_research: bool,
    db: Optional[Session],
) -> Tuple[Path, int]:
    """Generate TXT deliverable file."""
    lines = []

    # Header
    lines.append("=" * 60)
    lines.append("30-DAY CONTENT JUMPSTART DELIVERABLE")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"Client: {client.name}")
    lines.append(f"Project: {project.name}")
    lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"Total Posts: {len(posts)}")
    lines.append("")
    lines.append("=" * 60)
    lines.append("")

    # Posts
    for i, post in enumerate(posts, 1):
        lines.append(f"POST #{i}")
        lines.append("-" * 40)
        if post.template_name:
            lines.append(f"Template: {post.template_name}")
        if post.target_platform:
            lines.append(f"Platform: {post.target_platform}")
        lines.append(f"Word Count: {post.word_count or 'N/A'}")
        lines.append(f"Has CTA: {'Yes' if post.has_cta else 'No'}")
        if post.readability_score:
            lines.append(f"Readability: {post.readability_score:.1f}")
        lines.append("")
        lines.append(post.content)
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

    # Research results (if requested)
    if include_research and db:
        research_sections = _generate_research_section(project.id, db)
        if research_sections:
            lines.append("")
            lines.append("=" * 60)
            lines.append("RESEARCH RESULTS")
            lines.append("=" * 60)
            lines.append("")

            for tool_name, section_lines in research_sections.items():
                lines.extend(section_lines)
                lines.append("")

    # Audit log (if requested)
    if include_audit_log and db:
        audit_section = _generate_audit_section(project.id, db)
        if audit_section:
            lines.append("")
            lines.append("AUDIT LOG")
            lines.append("-" * 40)
            lines.extend(audit_section)

    # Write file
    content = "\n".join(lines)
    output_path.write_text(content, encoding="utf-8")

    file_size = output_path.stat().st_size
    logger.info(f"Generated TXT deliverable: {output_path} ({file_size} bytes)")

    return output_path, file_size


async def _generate_markdown(
    posts: List[Post],
    client: Client,
    project: Project,
    output_path: Path,
    include_audit_log: bool,
    include_research: bool,
    db: Optional[Session],
) -> Tuple[Path, int]:
    """Generate Markdown deliverable file with frontmatter."""
    lines = []

    # YAML frontmatter
    lines.append("---")
    lines.append("title: 30-Day Content Jumpstart Deliverable")
    lines.append(f"client: {client.name}")
    lines.append(f"project: {project.name}")
    lines.append(f'generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    lines.append(f"total_posts: {len(posts)}")
    lines.append("format: markdown")
    lines.append("---")
    lines.append("")

    # Header
    lines.append("# 30-Day Content Jumpstart Deliverable")
    lines.append("")
    lines.append(f"**Client:** {client.name}")
    lines.append(f"**Project:** {project.name}")
    lines.append(f'**Generated:** {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    lines.append(f"**Total Posts:** {len(posts)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Table of Contents
    lines.append("## Table of Contents")
    lines.append("")
    for i, post in enumerate(posts, 1):
        template_name = post.template_name or "Custom"
        platform = f" ({post.target_platform})" if post.target_platform else ""
        lines.append(f"{i}. [Post {i}: {template_name}{platform}](#post-{i})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Posts section
    lines.append("## Your Posts")
    lines.append("")

    for i, post in enumerate(posts, 1):
        # Post header with anchor
        template_name = post.template_name or "Custom"
        lines.append(f"### Post {i}: {template_name} {{#post-{i}}}")
        lines.append("")

        # Metadata table
        lines.append("| Property | Value |")
        lines.append("|----------|-------|")
        if post.target_platform:
            lines.append(f"| **Platform** | {post.target_platform} |")
        lines.append(f'| **Word Count** | {post.word_count or "N/A"} |')
        lines.append(f'| **Has CTA** | {"✅ Yes" if post.has_cta else "❌ No"} |')
        if post.readability_score:
            lines.append(f"| **Readability Score** | {post.readability_score:.1f} |")
        lines.append("")

        # Post content in blockquote for easy copying
        lines.append("#### Content")
        lines.append("")
        # Add content as blockquote (simpler approach - just prepend > to content)
        for line in post.content.split("\n"):
            lines.append(f"> {line}")
        lines.append("")

        # Separator
        lines.append("---")
        lines.append("")

    # Research results (if requested)
    if include_research and db:
        research_sections = _generate_research_section(project.id, db)
        if research_sections:
            lines.append("---")
            lines.append("")
            lines.append("## Research Results")
            lines.append("")

            for tool_name, section_lines in research_sections.items():
                lines.extend(section_lines)
                lines.append("")

    # Audit log (if requested)
    if include_audit_log and db:
        audit_section = _generate_audit_section(project.id, db)
        if audit_section:
            lines.append("## Audit Log")
            lines.append("")
            lines.append("```")
            lines.extend(audit_section)
            lines.append("```")
            lines.append("")

    # Footer
    lines.append("---")
    lines.append("")
    lines.append("*Generated by 30-Day Content Jumpstart*")
    lines.append("")

    # Write file
    content = "\n".join(lines)
    output_path.write_text(content, encoding="utf-8")

    file_size = output_path.stat().st_size
    logger.info(f"Generated Markdown deliverable: {output_path} ({file_size} bytes)")

    return output_path, file_size


async def _generate_docx(
    posts: List[Post],
    client: Client,
    project: Project,
    output_path: Path,
    include_audit_log: bool,
    include_research: bool,
    db: Optional[Session],
) -> Tuple[Path, int]:
    """Generate DOCX deliverable file."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed, falling back to TXT")
        # Fall back to TXT if docx not available
        txt_path = output_path.with_suffix(".txt")
        return await _generate_txt(posts, client, project, txt_path, include_audit_log, db)

    # Create document
    doc = Document()

    # Brand color
    brand_color = RGBColor(41, 128, 185)

    # Title page
    title = doc.add_paragraph()
    title_run = title.add_run("30-Day Content Jumpstart")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = brand_color
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Client name
    client_para = doc.add_paragraph()
    client_run = client_para.add_run(client.name)
    client_run.font.name = "Calibri"
    client_run.font.size = Pt(22)
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(12)
    date_run.font.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Introduction
    doc.add_heading("About This Content Package", level=1)
    doc.add_paragraph(
        f"This content package has been custom-generated for {client.name}. "
        f"All {len(posts)} posts are tailored to your brand voice, target audience, and business goals."
    )

    doc.add_paragraph()

    # Client info table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Client"
    table.rows[0].cells[1].text = client.name
    table.rows[1].cells[0].text = "Project"
    table.rows[1].cells[1].text = project.name
    table.rows[2].cells[0].text = "Total Posts"
    table.rows[2].cells[1].text = str(len(posts))

    doc.add_page_break()

    # Posts section
    doc.add_heading("Your Posts", level=1)

    for i, post in enumerate(posts, 1):
        # Post header
        doc.add_heading(f"Post {i}: {post.template_name or 'Custom'}", level=2)

        # Post content
        content_para = doc.add_paragraph(post.content)
        content_para.paragraph_format.space_after = Pt(12)

        # Metadata
        metadata_parts = [f"Words: {post.word_count or 'N/A'}"]
        if post.target_platform:
            metadata_parts.insert(0, f"Platform: {post.target_platform}")
        metadata_parts.append(f"Has CTA: {'Yes' if post.has_cta else 'No'}")
        if post.readability_score:
            metadata_parts.append(f"Readability: {post.readability_score:.1f}")

        metadata_para = doc.add_paragraph(" | ".join(metadata_parts))
        metadata_para.runs[0].font.size = Pt(9)
        metadata_para.runs[0].font.italic = True
        metadata_para.runs[0].font.color.rgb = RGBColor(127, 140, 141)

        # Separator
        doc.add_paragraph("_" * 80)

        # Page break every 5 posts
        if i % 5 == 0 and i < len(posts):
            doc.add_page_break()

    # Research results (if requested)
    if include_research and db:
        research_sections = _generate_research_section(project.id, db)
        if research_sections:
            doc.add_page_break()
            doc.add_heading("Research Results", level=1)

            for tool_name, section_lines in research_sections.items():
                for line in section_lines:
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=2)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=3)
                    elif line.startswith("|"):
                        # Table row - skip for now (complex DOCX table formatting)
                        doc.add_paragraph(line)
                    elif line:
                        doc.add_paragraph(line)

    # Audit log (if requested)
    if include_audit_log and db:
        audit_section = _generate_audit_section(project.id, db)
        if audit_section:
            doc.add_page_break()
            doc.add_heading("Audit Log", level=1)
            for line in audit_section:
                doc.add_paragraph(line)

    # Save document
    doc.save(str(output_path))

    file_size = output_path.stat().st_size
    logger.info(f"Generated DOCX deliverable: {output_path} ({file_size} bytes)")

    return output_path, file_size


def _generate_audit_section(project_id: str, db: Session) -> List[str]:
    """Generate audit log section for a project."""
    lines = []

    try:
        from backend.models import AuditLog

        # Query audit logs for this project
        audit_logs = (
            db.query(AuditLog)
            .filter(
                (AuditLog.target_id == project_id)
                | (AuditLog.metadata.contains({"project_id": project_id}))
            )
            .order_by(AuditLog.timestamp.asc())
            .limit(100)
            .all()
        )

        if not audit_logs:
            return []

        for log in audit_logs:
            timestamp = log.timestamp.strftime("%Y-%m-%d %H:%M:%S") if log.timestamp else "N/A"
            lines.append(f"[{timestamp}] {log.actor}: {log.action} on {log.target_type}")

    except Exception as e:
        logger.warning(f"Could not generate audit section: {e}")
        return []

    return lines


def _generate_research_section(project_id: str, db: Session) -> dict:
    """
    Generate research results section for deliverables.

    Args:
        project_id: Project ID
        db: Database session

    Returns:
        Dict mapping tool_name -> list of formatted lines
    """
    from backend.services import crud

    results = crud.get_research_results_by_project(db, project_id, status="completed")
    if not results:
        return {}

    research_sections = {}

    for result in results:
        lines = []
        lines.append(f"# {result.tool_label or result.tool_name}")
        lines.append("")

        # Metadata
        if result.created_at:
            lines.append(f"**Executed:** {result.created_at.strftime('%B %d, %Y at %H:%M')}")
        if result.tool_price:
            lines.append(f"**Cost:** ${result.tool_price:.2f}")
        lines.append("")

        # Format tool-specific data
        if result.data:
            lines.append("## Executive Summary")
            lines.append("")

            if result.tool_name == "voice_analysis":
                lines.extend(_format_voice_analysis(result.data))
            elif result.tool_name == "brand_archetype":
                lines.extend(_format_brand_archetype(result.data))
            elif result.tool_name == "seo_keyword_research":
                lines.extend(_format_seo_keywords(result.data))
            elif result.tool_name == "determine_competitors":
                lines.extend(_format_determine_competitors(result.data))
            elif result.tool_name == "competitive_analysis":
                lines.extend(_format_competitive_analysis(result.data))
            elif result.tool_name == "content_gap_analysis":
                lines.extend(_format_content_gap(result.data))
            elif result.tool_name == "market_trends_research":
                lines.extend(_format_market_trends(result.data))
            else:
                # Generic fallback for other tools
                lines.extend(_format_generic_research(result.data))

            lines.append("")

        # Full report from output files
        if result.outputs:
            full_report_lines = _read_output_files(result.outputs)
            if full_report_lines:
                lines.append("---")
                lines.append("")
                lines.append("## Full Report")
                lines.append("")
                lines.extend(full_report_lines)
                lines.append("")

        research_sections[result.tool_name] = lines

    return research_sections


def _read_output_files(outputs: dict) -> List[str]:
    """
    Read and format output files from research tools.

    Args:
        outputs: Dict mapping format -> file path
                 Example: {"markdown": "path/to/report.md"}

    Returns:
        List of formatted lines to include in deliverable
    """
    lines: List[str] = []

    if not outputs or not isinstance(outputs, dict):
        return lines

    # Prefer markdown format for readability
    if "markdown" in outputs and outputs["markdown"]:
        markdown_path = Path(outputs["markdown"])
        try:
            if markdown_path.exists() and markdown_path.is_file():
                content = markdown_path.read_text(encoding="utf-8")
                lines.append(content)
                logger.info(f"Included markdown output: {markdown_path}")
            else:
                logger.warning(f"Markdown output file not found: {markdown_path}")
                lines.append("*Full markdown report unavailable (file not found)*")
        except Exception as e:
            logger.error(f"Error reading markdown output {markdown_path}: {e}")
            lines.append(f"*Full report unavailable (error: {e})*")

    # Fallback to JSON if markdown not available
    elif "json" in outputs and outputs["json"]:
        json_path = Path(outputs["json"])
        try:
            if json_path.exists() and json_path.is_file():
                with json_path.open("r", encoding="utf-8") as f:
                    data = json.load(f)
                lines.append("```json")
                lines.append(json.dumps(data, indent=2, ensure_ascii=False))
                lines.append("```")
                logger.info(f"Included JSON output: {json_path}")
            else:
                logger.warning(f"JSON output file not found: {json_path}")
                lines.append("*Full JSON report unavailable (file not found)*")
        except Exception as e:
            logger.error(f"Error reading JSON output {json_path}: {e}")
            lines.append(f"*Full report unavailable (error: {e})*")
    else:
        logger.info("No markdown or JSON output available")
        lines.append("*Full report not available for this tool*")

    return lines


def _format_voice_analysis(data: dict) -> List[str]:
    """Format comprehensive voice analysis results."""
    lines = []

    # Summary
    if "summary" in data:
        lines.append(f"**Voice Summary:** {data['summary']}")
        lines.append("")

    # Tone & Personality
    lines.append("### Tone & Personality")
    lines.append("")

    primary_tone = data.get("primary_tone", data.get("tone", "Unknown"))
    secondary_tone = data.get("secondary_tone")
    tone_text = f"**Primary Tone:** {primary_tone}"
    if secondary_tone:
        tone_text += f" | **Secondary:** {secondary_tone}"
    lines.append(tone_text)

    if "formality_score" in data:
        formality = data["formality_score"]
        formality_label = "Formal" if formality >= 7 else "Casual" if formality <= 4 else "Balanced"
        lines.append(f"**Formality:** {formality_label} ({formality}/10)")

    if "confidence_score" in data:
        confidence = data["confidence_score"]
        confidence_label = (
            "Very Confident" if confidence >= 8 else "Confident" if confidence >= 6 else "Moderate"
        )
        lines.append(f"**Confidence Level:** {confidence_label} ({confidence}/10)")

    if "personality_traits" in data and data["personality_traits"]:
        traits = ", ".join(
            [str(t).replace("_", " ").title() for t in data["personality_traits"][:5]]
        )
        lines.append(f"**Key Traits:** {traits}")

    lines.append("")

    # Writing Patterns
    lines.append("### Writing Patterns")
    lines.append("")

    # Sentence structure
    if "sentence_analysis" in data:
        sa = data["sentence_analysis"]
        avg_len = sa.get("avg_length", 0)
        variety = sa.get("variety_score", 0)
        lines.append(
            f"**Sentence Length:** {avg_len:.1f} words average (variety: {variety:.1f}/10)"
        )

    # Punctuation style
    if "punctuation_analysis" in data:
        pa = data["punctuation_analysis"]
        punct_style = []
        if pa.get("exclamation_marks", 0) > 2:
            punct_style.append("uses exclamations")
        if pa.get("questions", 0) > 2:
            punct_style.append("asks questions")
        if pa.get("em_dashes", 0) > 1:
            punct_style.append("uses em-dashes")
        if punct_style:
            lines.append(f"**Punctuation Style:** {', '.join(punct_style)}")

    # Vocabulary
    if "vocabulary_analysis" in data:
        va = data["vocabulary_analysis"]
        if "lexical_diversity" in va:
            diversity = va["lexical_diversity"]
            diversity_label = (
                "Rich" if diversity >= 0.6 else "Moderate" if diversity >= 0.4 else "Simple"
            )
            lines.append(f"**Vocabulary:** {diversity_label} (diversity: {diversity:.2f})")

    # Pronoun focus
    if "pronoun_focus" in data:
        focus = data["pronoun_focus"]
        focus_meaning = {"you": "Reader-focused", "we": "Collaborative", "I": "Personal/Individual"}
        lines.append(f"**Voice Focus:** {focus_meaning.get(focus, focus)}")

    lines.append("")

    # Content Patterns
    lines.append("### Content Patterns")
    lines.append("")

    # Opening patterns
    if "opening_patterns" in data and data["opening_patterns"]:
        lines.append("**Common Openings:**")
        for pattern in data["opening_patterns"][:3]:
            pattern_type = pattern.get("pattern_type", "")
            example = pattern.get("example", "")[:80]
            lines.append(f'- {pattern_type.title()}: "{example}..."')
        lines.append("")

    # Signature phrases
    if "signature_phrases" in data and data["signature_phrases"]:
        lines.append("**Signature Phrases:**")
        for phrase in data["signature_phrases"][:5]:
            phrase_text = phrase.get("phrase", "")
            count = phrase.get("count", 0)
            lines.append(f'- "{phrase_text}" (used {count}x)')
        lines.append("")

    # CTA patterns
    if "cta_patterns" in data and data["cta_patterns"]:
        lines.append("**CTA Style:**")
        for cta in data["cta_patterns"][:3]:
            cta_type = cta.get("cta_type", "")
            example = cta.get("example", "")[:60]
            lines.append(f'- {cta_type.title()}: "{example}"')
        lines.append("")

    # Data usage
    if "uses_data" in data and data["uses_data"]:
        freq = data.get("data_frequency", "sometimes")
        lines.append(f"**Uses Data/Statistics:** {freq.title()}")
        if "data_examples" in data and data["data_examples"]:
            example = data["data_examples"][0][:80]
            lines.append(f'  Example: "{example}..."')
        lines.append("")

    # Structure preferences
    structure_prefs = []
    if data.get("uses_bullets"):
        structure_prefs.append("bullet points")
    if data.get("uses_emojis"):
        structure_prefs.append("emojis")
    if data.get("uses_formatting"):
        structure_prefs.append("bold/italic formatting")
    if structure_prefs:
        lines.append(f"**Formatting:** Uses {', '.join(structure_prefs)}")
        lines.append("")

    # Readability
    lines.append("### Readability")
    lines.append("")

    if "readability_score" in data:
        score = data["readability_score"]
        reading_level = data.get("reading_level", "Unknown")
        # Flesch Reading Ease scale: 90-100 = Very Easy, 60-70 = Standard, 0-30 = Very Difficult
        ease_label = (
            "Very Easy"
            if score >= 80
            else "Easy" if score >= 70 else "Standard" if score >= 60 else "Difficult"
        )
        lines.append(f"**Readability:** {ease_label} ({score:.1f}/100 Flesch score)")
        lines.append(f"**Reading Level:** {reading_level}")

    lines.append("")

    # Recommendations
    lines.append("### Voice Quality")
    lines.append("")

    if "consistency_score" in data:
        consistency = data["consistency_score"]
        consistency_label = (
            "Excellent" if consistency >= 8 else "Good" if consistency >= 6 else "Needs Work"
        )
        lines.append(f"**Consistency:** {consistency_label} ({consistency:.1f}/10)")
        lines.append("")

    if "strengths" in data and data["strengths"]:
        lines.append("**Strengths:**")
        for strength in data["strengths"][:5]:
            lines.append(f"- {strength}")
        lines.append("")

    if "improvement_areas" in data and data["improvement_areas"]:
        lines.append("**Areas to Improve:**")
        for area in data["improvement_areas"][:5]:
            lines.append(f"- {area}")
        lines.append("")

    # Legacy recommendations field (backwards compatibility)
    if "recommendations" in data and isinstance(data["recommendations"], list):
        if not data.get("improvement_areas"):  # Only show if no improvement_areas
            lines.append("**Recommendations:**")
            for rec in data["recommendations"][:5]:
                lines.append(f"- {rec}")

    return lines


def _format_brand_archetype(data: dict) -> List[str]:
    """Format brand archetype results with comprehensive details.

    Displays archetype summary, confidence scores, trait matches, voice guidance,
    content themes, brand examples, and actionable recommendations.
    """
    lines = []

    # Get archetype details from the 12 archetypes
    # Import here to avoid circular dependency
    from src.research.brand_archetype import ARCHETYPES

    primary_id = data.get("primary_archetype")
    secondary_id = data.get("secondary_archetype")

    if not primary_id:
        lines.append("**Brand Archetype:** Not determined")
        return lines

    # Get full archetype objects
    primary_arch = ARCHETYPES.get(primary_id)
    secondary_arch = ARCHETYPES.get(secondary_id) if secondary_id else None

    if not primary_arch:
        lines.append(f"**Brand Archetype:** {primary_id}")
        return lines

    # 1. Archetype Summary
    lines.append("### Brand Archetype Summary")
    lines.append("")
    lines.append(f"**Primary Archetype:** {primary_arch.name}")

    # Confidence score
    confidence = data.get("confidence_score")
    if confidence is not None:
        lines.append(f"**Confidence Score:** {confidence:.1%}")

    lines.append("")
    lines.append(f"_{primary_arch.description}_")
    lines.append("")

    # Secondary archetype
    if secondary_arch:
        lines.append(f"**Secondary Archetype:** {secondary_arch.name}")
        lines.append(f"_{secondary_arch.description}_")
        lines.append("")

    # 2. Archetype Fit Scores (top competitors)
    archetype_scores = data.get("archetype_scores", {})
    if archetype_scores:
        lines.append("### Archetype Fit Analysis")
        lines.append("")
        lines.append("How well your brand fits each archetype:")
        lines.append("")
        for arch_id, score in list(archetype_scores.items())[:5]:  # Top 5
            arch = ARCHETYPES.get(arch_id)
            if arch:
                lines.append(f"- **{arch.name}:** {score:.1%}")
        lines.append("")

    # 3. Key Traits
    lines.append("### Key Brand Traits")
    lines.append("")

    trait_matches = data.get("trait_matches", {})
    if trait_matches and "primary" in trait_matches:
        matched = trait_matches["primary"].get("matched_traits", [])
        all_traits = primary_arch.traits

        if matched:
            lines.append(f"**Detected traits:** {', '.join(matched)}")
        lines.append(f"**Core {primary_arch.name} traits:** {', '.join(all_traits)}")
    else:
        # Fallback to archetype traits
        lines.append(f"**Core traits:** {', '.join(primary_arch.traits)}")

    lines.append("")

    # 4. Voice Characteristics
    lines.append("### Voice & Tone Guidance")
    lines.append("")
    lines.append("Your brand voice should be:")
    for char in primary_arch.voice_characteristics:
        lines.append(f"- {char}")

    if secondary_arch:
        lines.append("")
        lines.append(f"_Blend in {secondary_arch.name.lower()} elements:_")
        for char in secondary_arch.voice_characteristics[:2]:
            lines.append(f"- {char}")

    lines.append("")

    # 5. Content Themes
    lines.append("### Content Themes")
    lines.append("")
    lines.append("Focus your content on:")
    for theme in primary_arch.content_themes:
        lines.append(f"- {theme}")
    lines.append("")

    # 6. Brand Examples
    lines.append("### Brands to Study")
    lines.append("")
    lines.append(f"Brands that share the {primary_arch.name} archetype:")
    for example in primary_arch.examples:
        lines.append(f"- {example}")
    lines.append("")

    # 7. Recommendations
    recommendations = data.get("recommendations", [])
    if recommendations:
        lines.append("### Recommendations")
        lines.append("")
        for rec in recommendations:
            # Clean markdown formatting for consistent display
            lines.append(rec)
            lines.append("")

    return lines


def _format_seo_keywords(data: dict) -> List[str]:
    """Format SEO keyword results with comprehensive strategy details.

    Displays strategy summary, primary/secondary keywords with trends data,
    keyword clusters, quick wins, competitor gaps, and content priorities.
    """
    lines = []

    # 1. Strategy Summary (executive overview)
    if data.get("strategy_summary"):
        lines.append("### SEO Strategy Summary")
        lines.append("")
        lines.append(data["strategy_summary"])
        lines.append("")

    # 2. Primary Keywords (comprehensive table)
    primary_keywords = data.get("primary_keywords", [])
    if primary_keywords:
        lines.append("### Primary Keywords")
        lines.append("")
        lines.append(
            "| # | Keyword | Intent | Difficulty | Volume | Relevance | Quality | Trends |"
        )
        lines.append(
            "|---|---------|--------|------------|--------|-----------|---------|--------|"
        )

        for i, kw in enumerate(primary_keywords[:10], 1):
            keyword_name = kw.get("keyword", "N/A")
            intent = kw.get("search_intent", "N/A")
            difficulty = kw.get("difficulty", "N/A").upper()
            volume = kw.get("monthly_volume_estimate", "Unknown")
            relevance = kw.get("relevance_score", 0)
            quality = kw.get("quality_score")

            # Format quality score with badge
            quality_str = f"{quality:.0f}" if quality else "—"
            if quality and quality >= 80:
                quality_str += " ⭐"
            elif quality and quality >= 70:
                quality_str += " ✓"

            # Format trends data
            trend_score = kw.get("trend_score")
            trend_dir = kw.get("trend_direction")
            trends_str = "—"
            if trend_score is not None:
                trends_str = f"{trend_score:.0f}"
                if trend_dir == "rising":
                    trends_str += " ↗"
                elif trend_dir == "declining":
                    trends_str += " ↘"
                elif trend_dir == "seasonal":
                    trends_str += " ~"

            lines.append(
                f"| {i} | {keyword_name} | {intent} | {difficulty} | {volume} | "
                f"{relevance}/10 | {quality_str} | {trends_str} |"
            )

        lines.append("")

        # Show detailed info for top 3 keywords
        lines.append("**Top Keyword Details:**")
        lines.append("")
        for i, kw in enumerate(primary_keywords[:3], 1):
            lines.append(f"**{i}. {kw.get('keyword', 'N/A')}**")

            details = []
            if kw.get("long_tail"):
                details.append("Long-tail")
            if kw.get("question_based"):
                details.append("Question-based")
            if kw.get("seasonal"):
                details.append("Seasonal")

            if details:
                lines.append(f"- Type: {', '.join(details)}")

            related_topics = kw.get("related_topics", [])
            if related_topics:
                lines.append(f"- Related: {', '.join(related_topics[:3])}")

            related_queries = kw.get("related_queries", [])
            if related_queries:
                lines.append(f"- Related searches: {', '.join(related_queries[:3])}")

            lines.append("")

    # 3. Quick Win Keywords
    quick_wins = data.get("quick_win_keywords", [])
    if quick_wins:
        lines.append("### Quick Win Keywords")
        lines.append("")
        lines.append("Low-difficulty, high-relevance opportunities for immediate targeting:")
        lines.append("")
        for qw in quick_wins[:5]:
            lines.append(f"- {qw}")
        lines.append("")

    # 4. Keyword Clusters
    clusters = data.get("keyword_clusters", [])
    if clusters:
        lines.append("### Keyword Clusters")
        lines.append("")

        for cluster in clusters[:5]:  # Top 5 clusters
            theme = cluster.get("theme", "Unknown")
            priority = cluster.get("priority", "Medium").upper()
            primary_kw = cluster.get("primary_keyword", "")

            lines.append(f"**[{priority}] {theme}**")
            lines.append(f"- Primary: {primary_kw}")

            secondary = cluster.get("secondary_keywords", [])
            if secondary:
                lines.append(f"- Secondary: {', '.join(secondary[:5])}")

            suggestions = cluster.get("content_suggestions", [])
            if suggestions:
                lines.append(f"- Content idea: {suggestions[0]}")

            lines.append("")

    # 5. Content Priorities
    priorities = data.get("content_priorities", [])
    if priorities:
        lines.append("### Content Priorities")
        lines.append("")
        lines.append("Recommended content to create (in priority order):")
        lines.append("")
        for i, priority in enumerate(priorities, 1):
            lines.append(f"{i}. {priority}")
        lines.append("")

    # 6. Competitor Analysis (if available)
    competitor_analysis = data.get("competitor_analysis", [])
    if competitor_analysis:
        lines.append("### Competitor Insights")
        lines.append("")

        for comp in competitor_analysis[:3]:  # Top 3 competitors
            comp_name = comp.get("competitor_name", "Unknown")
            gaps = comp.get("gaps", [])
            overlaps = comp.get("overlaps", [])

            lines.append(f"**{comp_name}**")

            if gaps:
                lines.append(f"- Gap opportunities: {', '.join(gaps[:3])}")
            if overlaps:
                lines.append(f"- Keyword overlaps: {', '.join(overlaps[:3])}")

            lines.append("")

    # 7. Secondary Keywords (summary)
    secondary_keywords = data.get("secondary_keywords", [])
    if secondary_keywords:
        lines.append("### Secondary Keywords")
        lines.append("")
        lines.append(
            f"**{len(secondary_keywords)} long-tail keywords** identified for supporting content:"
        )
        lines.append("")

        # Show first 10
        for i, kw in enumerate(secondary_keywords[:10], 1):
            keyword_name = kw.get("keyword", "N/A")
            difficulty = kw.get("difficulty", "N/A")
            lines.append(f"{i}. {keyword_name} ({difficulty})")

        if len(secondary_keywords) > 10:
            lines.append(f"\n_... and {len(secondary_keywords) - 10} more_")

        lines.append("")

    # If no data available
    if not any([primary_keywords, quick_wins, clusters, priorities, secondary_keywords]):
        lines.append("**SEO Keywords:** No data available")

    return lines


def _format_competitive_analysis(data: dict) -> List[str]:
    """Format competitive analysis results with comprehensive competitor insights.

    Displays market landscape, detailed competitor profiles, content gaps, differentiation
    strategies, positioning recommendations, and actionable priorities.
    """
    lines = []

    if not data:
        lines.append("**Competitive Analysis:** No data available")
        return lines

    # 1. Market Landscape
    lines.append("### Competitive Landscape Overview")
    lines.append("")

    market_summary = data.get("market_summary", "")
    if market_summary:
        lines.append(market_summary)
        lines.append("")

    market_saturation = data.get("market_saturation", "")
    if market_saturation:
        lines.append(f"**Market Saturation:** {market_saturation}")
        lines.append("")

    # 2. Competitor Profiles (Detailed)
    competitors = data.get("competitors", [])
    if competitors:
        lines.append("### Competitor Profiles")
        lines.append("")

        for i, comp in enumerate(competitors[:5], 1):  # Top 5 competitors
            if isinstance(comp, dict):
                name = comp.get("name", "Unknown")
                lines.append(f"**{i}. {name}**")

                # Positioning
                positioning = comp.get("positioning", "")
                if positioning:
                    lines.append(f"**Positioning:** {positioning}")

                # Target audience
                target_audience = comp.get("target_audience", "")
                if target_audience:
                    lines.append(f"**Target Audience:** {target_audience}")

                # Content strategy
                content_types = comp.get("content_types", [])
                content_freq = comp.get("content_frequency", "")
                content_topics = comp.get("content_topics", [])

                if content_types or content_freq or content_topics:
                    lines.append("")
                    lines.append("*Content Strategy:*")

                    if content_types:
                        types_str = ", ".join(
                            [str(t).replace("_", " ").title() for t in content_types[:5]]
                        )
                        lines.append(f"  - **Formats:** {types_str}")

                    if content_freq:
                        lines.append(f"  - **Frequency:** {content_freq}")

                    if content_topics:
                        topics_str = ", ".join(content_topics[:5])
                        lines.append(f"  - **Topics:** {topics_str}")

                # Brand voice
                brand_voice = comp.get("brand_voice", "")
                tone_descriptors = comp.get("tone_descriptors", [])

                if brand_voice or tone_descriptors:
                    lines.append("")
                    lines.append("*Voice & Tone:*")

                    if brand_voice:
                        lines.append(f"  - {brand_voice}")

                    if tone_descriptors:
                        tones_str = ", ".join(tone_descriptors[:5])
                        lines.append(f"  - Tone: {tones_str}")

                # Strengths and weaknesses
                strengths = comp.get("strengths", [])
                weaknesses = comp.get("weaknesses", [])

                if strengths:
                    lines.append("")
                    lines.append("*Strengths:*")
                    for strength in strengths[:5]:
                        lines.append(f"  - {strength}")

                if weaknesses:
                    lines.append("")
                    lines.append("*Weaknesses (Your Opportunity):*")
                    for weakness in weaknesses[:5]:
                        lines.append(f"  - {weakness}")

                # Reach and engagement
                estimated_reach = comp.get("estimated_reach", "")
                engagement_level = comp.get("engagement_level", "")

                if estimated_reach or engagement_level:
                    lines.append("")
                    metrics = []
                    if estimated_reach:
                        metrics.append(f"Reach: {estimated_reach}")
                    if engagement_level:
                        eng_level = str(engagement_level).upper()
                        metrics.append(f"Engagement: {eng_level}")
                    lines.append(f"**Metrics:** {' | '.join(metrics)}")

                lines.append("")

    # 3. Content Gap Opportunities
    content_gaps = data.get("content_gaps", [])
    if content_gaps:
        lines.append("### Content Gap Opportunities")
        lines.append("")
        lines.append("Topics where competitors are weak - your opportunity to dominate:")
        lines.append("")

        for i, gap in enumerate(content_gaps[:5], 1):
            if isinstance(gap, dict):
                topic = gap.get("topic", "Unknown")
                description = gap.get("description", "")
                opp_score = gap.get("opportunity_score")
                competitors_missing = gap.get("competitors_missing", [])
                suggested_content = gap.get("suggested_content", [])

                # Opportunity score indicator
                score_indicator = ""
                if opp_score is not None:
                    if opp_score >= 8:
                        score_indicator = " [HIGH OPPORTUNITY]"
                    elif opp_score >= 6:
                        score_indicator = " [MEDIUM OPPORTUNITY]"

                lines.append(f"**{i}. {topic}**{score_indicator}")

                if description:
                    lines.append(f"   {description}")

                if opp_score is not None:
                    lines.append(f"   - **Opportunity Score:** {opp_score:.1f}/10")

                if competitors_missing:
                    missing_str = ", ".join(competitors_missing[:3])
                    lines.append(f"   - **Competitors Missing:** {missing_str}")

                if suggested_content:
                    lines.append("   - **Content Ideas:**")
                    for idea in suggested_content[:3]:
                        lines.append(f"     - {idea}")

                lines.append("")

    # 4. Quick Wins
    quick_wins = data.get("quick_wins", [])
    if quick_wins:
        lines.append("### Quick Win Opportunities")
        lines.append("")
        lines.append("Immediate actions to gain competitive advantage:")
        lines.append("")
        for i, win in enumerate(quick_wins[:5], 1):
            lines.append(f"{i}. {win}")
        lines.append("")

    # 5. Differentiation Strategies
    diff_strategies = data.get("differentiation_strategies", [])
    if diff_strategies:
        lines.append("### Differentiation Strategies")
        lines.append("")

        for strat in diff_strategies[:5]:
            if isinstance(strat, dict):
                strategy_name = strat.get("strategy_name", "Unknown")
                description = strat.get("description", "")
                difficulty = strat.get("difficulty", "")
                impact = strat.get("potential_impact", "")
                examples = strat.get("examples", [])

                lines.append(f"**{strategy_name}**")

                if description:
                    lines.append(f"   {description}")

                if difficulty or impact:
                    meta = []
                    if difficulty:
                        meta.append(f"Difficulty: {difficulty}")
                    if impact:
                        meta.append(f"Impact: {impact}")
                    lines.append(f"   - {' | '.join(meta)}")

                if examples:
                    lines.append("   - **Examples:**")
                    for example in examples[:3]:
                        lines.append(f"     - {example}")

                lines.append("")

    # 6. Recommended Market Positioning
    recommended_position = data.get("recommended_position", {})
    if recommended_position and isinstance(recommended_position, dict):
        lines.append("### Recommended Market Positioning")
        lines.append("")

        positioning_statement = recommended_position.get("positioning_statement", "")
        if positioning_statement:
            lines.append("**Positioning Statement:**")
            lines.append(f"{positioning_statement}")
            lines.append("")

        unique_angles = recommended_position.get("unique_angles", [])
        if unique_angles:
            lines.append("**Unique Angles to Emphasize:**")
            for angle in unique_angles[:5]:
                lines.append(f"- {angle}")
            lines.append("")

        competitive_advantages = recommended_position.get("competitive_advantages", [])
        if competitive_advantages:
            lines.append("**Your Competitive Advantages:**")
            for adv in competitive_advantages[:5]:
                lines.append(f"- {adv}")
            lines.append("")

        areas_to_improve = recommended_position.get("areas_to_improve", [])
        if areas_to_improve:
            lines.append("**Areas to Improve:**")
            for area in areas_to_improve[:5]:
                lines.append(f"- {area}")
            lines.append("")

    # 7. Priority Actions
    priority_actions = data.get("priority_actions", [])
    if priority_actions:
        lines.append("### Priority Actions")
        lines.append("")
        lines.append("Top actions to take based on competitive analysis:")
        lines.append("")
        for i, action in enumerate(priority_actions[:5], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 8. Competitive Threats
    competitive_threats = data.get("competitive_threats", [])
    if competitive_threats:
        lines.append("### Competitive Threats to Watch")
        lines.append("")
        for threat in competitive_threats[:5]:
            lines.append(f"- {threat}")
        lines.append("")

    return lines


def _format_content_gap(data: dict) -> List[str]:
    """Format content gap analysis results with comprehensive gap breakdown.

    Displays executive summary, prioritized gaps with full details, competitor analysis,
    buyer journey gaps, format gaps, quick wins, roadmap, and actionable recommendations.
    """
    lines = []

    if not data:
        lines.append("**Content Gap Analysis:** No data available")
        return lines

    # 1. Executive Summary
    lines.append("### Content Gap Analysis Summary")
    lines.append("")

    if "executive_summary" in data:
        lines.append(data["executive_summary"])
        lines.append("")

    # Total gaps and opportunity
    if "total_gaps_identified" in data:
        lines.append(f"**Total Gaps Identified:** {data['total_gaps_identified']}")
    if "estimated_opportunity" in data:
        lines.append(f"**Estimated Opportunity:** {data['estimated_opportunity']}")

    lines.append("")

    # 2. Critical Priority Gaps (Full Details)
    critical_gaps = data.get("critical_gaps", [])
    if critical_gaps:
        lines.append("### CRITICAL Priority Gaps")
        lines.append("")
        lines.append("Must-create content with highest business impact:")
        lines.append("")

        for i, gap in enumerate(critical_gaps[:5], 1):
            if isinstance(gap, dict):
                title = gap.get("gap_title", "Unknown")
                gap_type = gap.get("gap_type", "N/A").upper()
                lines.append(f"**{i}. {title}** ({gap_type})")

                # Description
                description = gap.get("description", "")
                if description:
                    lines.append(f"   {description}")

                # Search & Competition
                search_vol = gap.get("search_volume", "")
                competition = gap.get("competition", "")
                if search_vol or competition:
                    metrics = []
                    if search_vol:
                        metrics.append(f"Search: {search_vol}")
                    if competition:
                        metrics.append(f"Competition: {competition}")
                    lines.append(f"   - {' | '.join(metrics)}")

                # Impact & Audience
                impact = gap.get("business_impact", "")
                audience = gap.get("target_audience", "")
                if impact:
                    lines.append(f"   - **Impact:** {impact}")
                if audience:
                    lines.append(f"   - **Audience:** {audience}")

                # Buyer stage
                buyer_stage = gap.get("buyer_stage", "")
                if buyer_stage:
                    lines.append(f"   - **Buyer Stage:** {buyer_stage}")

                # Content angle
                content_angle = gap.get("content_angle", "")
                if content_angle:
                    lines.append(f"   - **Recommended Angle:** {content_angle}")

                # Example topics
                example_topics = gap.get("example_topics", [])
                if example_topics:
                    lines.append("   - **Example Topics:**")
                    for topic in example_topics[:3]:
                        lines.append(f"     - {topic}")

                # Effort estimate
                effort = gap.get("estimated_effort", "")
                if effort:
                    lines.append(f"   - **Effort:** {effort}")

                lines.append("")

    # 3. High Priority Gaps
    high_gaps = data.get("high_priority_gaps", [])
    if high_gaps:
        lines.append("### HIGH Priority Gaps")
        lines.append("")
        lines.append("Strong opportunities with good ROI:")
        lines.append("")

        for i, gap in enumerate(high_gaps[:5], 1):
            if isinstance(gap, dict):
                title = gap.get("gap_title", "Unknown")
                description = gap.get("description", "")
                search_vol = gap.get("search_volume", "")
                impact = gap.get("business_impact", "")

                lines.append(f"**{i}. {title}**")
                if description:
                    lines.append(f"   {description}")
                if search_vol:
                    lines.append(f"   - Search: {search_vol}")
                if impact:
                    lines.append(f"   - Impact: {impact}")
                lines.append("")

    # 4. Medium Priority Gaps
    medium_gaps = data.get("medium_priority_gaps", [])
    if medium_gaps:
        lines.append("### MEDIUM Priority Gaps")
        lines.append("")

        for i, gap in enumerate(medium_gaps[:5], 1):
            if isinstance(gap, dict):
                title = gap.get("gap_title", "Unknown")
                description = gap.get("description", "")
                lines.append(f"{i}. **{title}** - {description}")

        lines.append("")

    # 5. Quick Wins
    quick_wins = data.get("quick_wins", [])
    if quick_wins:
        lines.append("### Quick Wins")
        lines.append("")
        lines.append("Easy content to create with high impact:")
        lines.append("")
        for win in quick_wins[:7]:
            lines.append(f"- {win}")
        lines.append("")

    # 6. Buyer Journey Gaps
    buyer_journey_gaps = data.get("buyer_journey_gaps", [])
    if buyer_journey_gaps:
        lines.append("### Buyer Journey Gaps")
        lines.append("")

        for journey_gap in buyer_journey_gaps:
            if isinstance(journey_gap, dict):
                stage = journey_gap.get("stage", "Unknown")
                current = journey_gap.get("current_coverage", "")
                gap_desc = journey_gap.get("gap_description", "")
                recommended = journey_gap.get("recommended_content", [])
                priority = journey_gap.get("priority", "").upper()

                lines.append(f"**{stage} Stage** [{priority} PRIORITY]")
                if current:
                    lines.append(f"Current Coverage: {current}")
                if gap_desc:
                    lines.append(f"Gap: {gap_desc}")

                if recommended:
                    lines.append("Recommended Content:")
                    for rec in recommended[:5]:
                        if isinstance(rec, dict):
                            # Handle structured recommendations
                            rec_title = rec.get("title", rec.get("topic", str(rec)))
                            lines.append(f"  - {rec_title}")
                        else:
                            lines.append(f"  - {rec}")

                lines.append("")

    # 7. Format Gaps
    format_gaps = data.get("format_gaps", [])
    if format_gaps:
        lines.append("### Missing Content Formats")
        lines.append("")

        for format_gap in format_gaps:
            if isinstance(format_gap, dict):
                format_name = format_gap.get("format_name", "Unknown")
                why_needed = format_gap.get("why_needed", "")
                topics = format_gap.get("topics_to_cover", [])
                impact = format_gap.get("estimated_impact", "")

                lines.append(f"**{format_name}** ({impact} impact)")
                if why_needed:
                    lines.append(f"  {why_needed}")

                if topics:
                    lines.append("  Topics to cover:")
                    for topic in topics[:3]:
                        lines.append(f"    - {topic}")

                lines.append("")

    # 8. Competitor Analysis
    competitor_analysis = data.get("competitor_analysis", [])
    if competitor_analysis:
        lines.append("### Competitor Content Analysis")
        lines.append("")

        for comp in competitor_analysis:
            if isinstance(comp, dict):
                name = comp.get("competitor_name", "Unknown")
                strengths = comp.get("content_strengths", [])
                popular = comp.get("popular_topics", [])
                formats = comp.get("formats_used", [])
                their_gaps = comp.get("gaps_in_their_content", [])

                lines.append(f"**{name}**")

                if strengths:
                    lines.append(f"  Strengths: {', '.join(strengths[:3])}")

                if popular:
                    lines.append(f"  Popular Topics: {', '.join(popular[:3])}")

                if formats:
                    lines.append(f"  Formats: {', '.join(formats[:3])}")

                if their_gaps:
                    lines.append(f"  Their Gaps (Your Opportunity): {', '.join(their_gaps[:3])}")

                lines.append("")

    # 9. Immediate Actions
    actions = data.get("immediate_actions", [])
    if actions:
        lines.append("### Immediate Actions")
        lines.append("")
        lines.append("What to do first:")
        lines.append("")
        for i, action in enumerate(actions[:5], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 10. Long-Term Opportunities
    long_term = data.get("long_term_opportunities", [])
    if long_term:
        lines.append("### Long-Term Strategic Opportunities")
        lines.append("")
        for opp in long_term[:5]:
            lines.append(f"- {opp}")
        lines.append("")

    # 11. 90-Day Roadmap
    roadmap = data.get("ninety_day_roadmap", [])
    if roadmap:
        lines.append("### 90-Day Content Roadmap")
        lines.append("")
        for i, item in enumerate(roadmap[:12], 1):
            lines.append(f"{i}. {item}")
        lines.append("")

    return lines


def _format_content_audit(data: dict) -> List[str]:
    """Format content audit results with comprehensive analysis.

    Displays executive summary, content inventory, performance analysis, topic analysis,
    refresh/repurpose/archive opportunities, strategic insights, and action plans.
    """
    lines = []

    if not data:
        lines.append("**Content Audit:** No data available")
        return lines

    # 1. Executive Summary
    lines.append("### Content Audit Executive Summary")
    lines.append("")

    total_pieces = data.get("total_content_pieces", 0)
    health_score = data.get("overall_health_score")

    if total_pieces:
        lines.append(f"**Total Content Analyzed:** {total_pieces} pieces")

    if health_score is not None:
        health_label = (
            "Excellent"
            if health_score >= 80
            else "Good" if health_score >= 60 else "Needs Improvement"
        )
        lines.append(f"**Overall Health Score:** {health_score:.1f}/100 ({health_label})")

    lines.append("")

    executive_summary = data.get("executive_summary", "")
    if executive_summary:
        lines.append(executive_summary)
        lines.append("")

    # 2. Content Distribution Metrics
    content_by_type = data.get("content_by_type", {})
    content_by_health = data.get("content_by_health", {})
    content_by_performance = data.get("content_by_performance", {})

    if content_by_type or content_by_health or content_by_performance:
        lines.append("### Content Distribution")
        lines.append("")

        if content_by_type:
            lines.append("**By Content Type:**")
            for content_type, count in sorted(
                content_by_type.items(), key=lambda x: x[1], reverse=True
            ):
                type_label = str(content_type).replace("_", " ").title()
                lines.append(f"  - {type_label}: {count}")
            lines.append("")

        if content_by_health:
            lines.append("**By Health Status:**")
            for health, count in sorted(
                content_by_health.items(), key=lambda x: x[1], reverse=True
            ):
                health_label = str(health).replace("_", " ").title()
                lines.append(f"  - {health_label}: {count}")
            lines.append("")

        if content_by_performance:
            lines.append("**By Performance Level:**")
            for perf, count in sorted(
                content_by_performance.items(), key=lambda x: x[1], reverse=True
            ):
                perf_label = str(perf).replace("_", " ").title()
                lines.append(f"  - {perf_label}: {count}")
            lines.append("")

    # 3. Top Performers
    top_performers = data.get("top_performers", [])
    if top_performers:
        lines.append("### Top Performing Content")
        lines.append("")
        lines.append("Your best content - double down on what works:")
        lines.append("")

        for i, piece in enumerate(top_performers[:10], 1):
            if isinstance(piece, dict):
                title = piece.get("title", "Unknown")
                perf_level = piece.get("performance_level", "").replace("_", " ").title()
                engagement = piece.get("engagement_score")
                traffic = piece.get("traffic_estimate", "")

                lines.append(f"**{i}. {title}** ({perf_level})")

                metrics = []
                if engagement is not None:
                    metrics.append(f"Engagement: {engagement:.0f}/100")
                if traffic:
                    metrics.append(f"Traffic: {traffic}")

                if metrics:
                    lines.append(f"   - {' | '.join(metrics)}")

                strengths = piece.get("strengths", [])
                if strengths:
                    lines.append(f"   - **Why It Works:** {', '.join(strengths[:3])}")

                lines.append("")

    # 4. Underperformers
    underperformers = data.get("underperformers", [])
    if underperformers:
        lines.append("### Underperforming Content")
        lines.append("")
        lines.append("Content needing attention:")
        lines.append("")

        for i, piece in enumerate(underperformers[:10], 1):
            if isinstance(piece, dict):
                title = piece.get("title", "Unknown")
                health = piece.get("health_status", "").replace("_", " ").title()
                action = piece.get("recommended_action", "")
                priority = piece.get("action_priority", "")

                priority_label = f" [{priority.upper()} PRIORITY]" if priority else ""
                lines.append(f"**{i}. {title}** ({health}){priority_label}")

                if action:
                    lines.append(f"   - **Action:** {action}")

                weaknesses = piece.get("weaknesses", [])
                if weaknesses:
                    lines.append(f"   - **Issues:** {', '.join(weaknesses[:3])}")

                updates_needed = piece.get("specific_updates_needed", [])
                if updates_needed:
                    lines.append("   - **Updates Needed:**")
                    for update in updates_needed[:3]:
                        lines.append(f"     - {update}")

                lines.append("")

    # 5. Topic Performance Analysis
    topic_performance = data.get("topic_performance", [])
    if topic_performance:
        lines.append("### Performance by Topic")
        lines.append("")

        for topic_data in topic_performance[:10]:
            if isinstance(topic_data, dict):
                topic = topic_data.get("topic", "Unknown")
                content_count = topic_data.get("content_count", 0)
                avg_perf = topic_data.get("avg_performance", "")
                top_piece = topic_data.get("top_performing_piece", "")
                recommendation = topic_data.get("recommendation", "")

                lines.append(f"**{topic}** ({content_count} pieces | Avg: {avg_perf})")

                if top_piece:
                    lines.append(f"  - **Best:** {top_piece}")

                underperf = topic_data.get("underperforming_pieces", [])
                if underperf:
                    lines.append(f"  - **Needs Work:** {', '.join(underperf[:2])}")

                if recommendation:
                    lines.append(f"  - **Strategy:** {recommendation}")

                lines.append("")

    # 6. Content Gaps
    content_gaps = data.get("content_gaps", [])
    if content_gaps:
        lines.append("### Content Gaps & Opportunities")
        lines.append("")

        for i, gap in enumerate(content_gaps[:10], 1):
            if isinstance(gap, dict):
                gap_desc = gap.get("gap_description", "Unknown")
                content_type = gap.get("content_type_needed", "")
                priority = gap.get("priority", "")
                reason = gap.get("reason", "")

                priority_label = f" [{priority.upper()}]" if priority else ""
                lines.append(f"**{i}. {gap_desc}**{priority_label}")

                if content_type:
                    lines.append(f"   - **Create:** {content_type}")

                if reason:
                    lines.append(f"   - **Why:** {reason}")

                lines.append("")

    # 7. Refresh Opportunities
    refresh_opps = data.get("refresh_opportunities", [])
    if refresh_opps:
        lines.append("### Content Refresh Opportunities")
        lines.append("")
        lines.append("Update existing content for quick wins:")
        lines.append("")

        for i, opp in enumerate(refresh_opps[:10], 1):
            if isinstance(opp, dict):
                content_title = opp.get("content_title", "Unknown")
                last_updated = opp.get("last_updated", "")
                why_refresh = opp.get("why_refresh", "")
                approach = opp.get("refresh_approach", "")
                impact = opp.get("estimated_impact", "")
                effort = opp.get("estimated_effort", "")

                impact_effort = []
                if impact:
                    impact_effort.append(f"Impact: {impact}")
                if effort:
                    impact_effort.append(f"Effort: {effort}")
                impact_label = f" ({' | '.join(impact_effort)})" if impact_effort else ""

                lines.append(f"**{i}. {content_title}**{impact_label}")

                if last_updated:
                    lines.append(f"   - **Last Updated:** {last_updated}")

                if why_refresh:
                    lines.append(f"   - **Why:** {why_refresh}")

                if approach:
                    lines.append(f"   - **How:** {approach}")

                lines.append("")

    # 8. Repurpose Opportunities
    repurpose_opps = data.get("repurpose_opportunities", [])
    if repurpose_opps:
        lines.append("### Content Repurposing Opportunities")
        lines.append("")
        lines.append("Get more mileage from existing content:")
        lines.append("")

        for i, opp in enumerate(repurpose_opps[:10], 1):
            if isinstance(opp, dict):
                source = opp.get("source_content", "Unknown")
                repurpose_into = opp.get("repurpose_into", "")
                platform = opp.get("target_platform", "")
                why = opp.get("why_repurpose", "")
                reach = opp.get("estimated_reach", "")

                lines.append(f"**{i}. {source}**")

                if repurpose_into:
                    lines.append(f"   - **Repurpose Into:** {repurpose_into}")

                if platform:
                    lines.append(f"   - **Platform:** {platform}")

                if why:
                    lines.append(f"   - **Value:** {why}")

                if reach:
                    lines.append(f"   - **Potential Reach:** {reach}")

                lines.append("")

    # 9. Archive Recommendations
    archive_recs = data.get("archive_recommendations", [])
    if archive_recs:
        lines.append("### Archive & Consolidation Recommendations")
        lines.append("")

        for i, rec in enumerate(archive_recs[:10], 1):
            if isinstance(rec, dict):
                content_title = rec.get("content_title", "Unknown")
                reason = rec.get("reason", "")
                action = rec.get("action", "")
                consolidate_into = rec.get("consolidate_into")

                lines.append(f"**{i}. {content_title}**")

                if action:
                    lines.append(f"   - **Action:** {action}")

                if consolidate_into:
                    lines.append(f"   - **Consolidate Into:** {consolidate_into}")

                if reason:
                    lines.append(f"   - **Reason:** {reason}")

                lines.append("")

    # 10. Strategic Insights
    content_strengths = data.get("content_strengths", [])
    content_weaknesses = data.get("content_weaknesses", [])

    if content_strengths or content_weaknesses:
        lines.append("### Strategic Insights")
        lines.append("")

        if content_strengths:
            lines.append("**What's Working Well:**")
            for strength in content_strengths[:5]:
                lines.append(f"- {strength}")
            lines.append("")

        if content_weaknesses:
            lines.append("**What Needs Improvement:**")
            for weakness in content_weaknesses[:5]:
                lines.append(f"- {weakness}")
            lines.append("")

    # 11. Immediate Actions
    immediate_actions = data.get("immediate_actions", [])
    if immediate_actions:
        lines.append("### Immediate Actions")
        lines.append("")
        lines.append("Priority actions for the next 7 days:")
        lines.append("")
        for i, action in enumerate(immediate_actions[:5], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 12. 30-Day Plan
    thirty_day_plan = data.get("thirty_day_plan", [])
    if thirty_day_plan:
        lines.append("### 30-Day Action Plan")
        lines.append("")
        for i, action in enumerate(thirty_day_plan[:10], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 13. 90-Day Plan
    ninety_day_plan = data.get("ninety_day_plan", [])
    if ninety_day_plan:
        lines.append("### 90-Day Strategic Plan")
        lines.append("")
        for i, action in enumerate(ninety_day_plan[:12], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    return lines


def _format_platform_strategy(data: dict) -> List[str]:
    """Format platform strategy results with comprehensive platform analysis.

    Displays executive summary, recommended platform mix, audience behavior, detailed
    platform recommendations, content distribution strategy, and implementation plans.
    """
    lines = []

    if not data:
        lines.append("**Platform Strategy:** No data available")
        return lines

    # 1. Executive Summary
    lines.append("### Platform Strategy Executive Summary")
    lines.append("")

    executive_summary = data.get("executive_summary", "")
    if executive_summary:
        lines.append(executive_summary)
        lines.append("")

    # 2. Recommended Platform Mix
    platform_mix = data.get("recommended_platform_mix", {})
    if platform_mix and isinstance(platform_mix, dict):
        lines.append("### Recommended Platform Mix")
        lines.append("")

        rationale = platform_mix.get("rationale", "")
        if rationale:
            lines.append(f"**Strategy:** {rationale}")
            lines.append("")

        primary = platform_mix.get("primary_platforms", [])
        if primary:
            primary_str = ", ".join([str(p).replace("_", " ").title() for p in primary])
            lines.append(f"**Primary Platforms (70% effort):** {primary_str}")
            lines.append("Focus your core content efforts here for maximum ROI.")
            lines.append("")

        secondary = platform_mix.get("secondary_platforms", [])
        if secondary:
            secondary_str = ", ".join([str(p).replace("_", " ").title() for p in secondary])
            lines.append(f"**Secondary Platforms (25% effort):** {secondary_str}")
            lines.append("Repurpose and cross-post primary content here.")
            lines.append("")

        experimental = platform_mix.get("experimental_platforms", [])
        if experimental:
            experimental_str = ", ".join([str(p).replace("_", " ").title() for p in experimental])
            lines.append(f"**Experimental Platforms (5% effort):** {experimental_str}")
            lines.append("Test and learn with minimal investment.")
            lines.append("")

        avoid = platform_mix.get("avoid_platforms", [])
        if avoid:
            avoid_str = ", ".join([str(p).replace("_", " ").title() for p in avoid])
            lines.append(f"**Platforms to Avoid (for now):** {avoid_str}")
            lines.append("")

    # 3. Audience Behavior Analysis
    audience_behavior = data.get("audience_behavior", [])
    if audience_behavior:
        lines.append("### Where Your Audience Is Active")
        lines.append("")

        for behavior in audience_behavior:
            if isinstance(behavior, dict):
                platform = str(behavior.get("platform", "Unknown")).replace("_", " ").title()
                audience_present = behavior.get("audience_present", False)
                activity_level = behavior.get("activity_level", "")
                consumption = behavior.get("content_consumption_pattern", "")
                engagement = behavior.get("engagement_style", "")
                decision_makers = behavior.get("decision_maker_presence", "")

                presence_indicator = "✓" if audience_present else "✗"
                lines.append(f"**{platform}** [{presence_indicator} Audience Present]")

                if activity_level:
                    lines.append(f"  - **Activity Level:** {activity_level}")

                if consumption:
                    lines.append(f"  - **Content Consumption:** {consumption}")

                if engagement:
                    lines.append(f"  - **Engagement Style:** {engagement}")

                if decision_makers:
                    lines.append(f"  - **Decision Makers:** {decision_makers}")

                lines.append("")

    # 4. Detailed Platform Recommendations
    platform_recs = data.get("platform_recommendations", [])
    if platform_recs:
        lines.append("### Detailed Platform Recommendations")
        lines.append("")

        for rec in platform_recs:
            if isinstance(rec, dict):
                platform = str(rec.get("platform", "Unknown")).replace("_", " ").title()
                fit_level = str(rec.get("fit_level", "")).replace("_", " ").title()
                priority = rec.get("priority", "").upper()

                # Platform header with fit level
                priority_label = f" [{priority} PRIORITY]" if priority else ""
                lines.append(f"**{platform}** ({fit_level}){priority_label}")
                lines.append("")

                # Why use / why not
                why_use = rec.get("why_use", [])
                why_not = rec.get("why_not_use", [])

                if why_use:
                    lines.append("**Why Use This Platform:**")
                    for reason in why_use[:5]:
                        lines.append(f"  ✓ {reason}")
                    lines.append("")

                if why_not:
                    lines.append("**Considerations:**")
                    for concern in why_not[:3]:
                        lines.append(f"  ⚠ {concern}")
                    lines.append("")

                # Strategy
                recommended_formats = rec.get("recommended_formats", [])
                posting_freq = rec.get("posting_frequency", "")
                content_approach = rec.get("content_approach", "")

                lines.append("**Content Strategy:**")

                if recommended_formats:
                    formats_str = ", ".join(
                        [str(f).replace("_", " ").title() for f in recommended_formats]
                    )
                    lines.append(f"  - **Formats:** {formats_str}")

                if posting_freq:
                    lines.append(f"  - **Frequency:** {posting_freq}")

                if content_approach:
                    lines.append(f"  - **Approach:** {content_approach}")

                lines.append("")

                # Expected outcomes
                primary_goal = rec.get("primary_goal", "")
                success_metrics = rec.get("success_metrics", [])
                effort = rec.get("estimated_effort", "")
                roi = rec.get("expected_ROI", "")

                if primary_goal or success_metrics or effort or roi:
                    lines.append("**Expected Outcomes:**")

                    if primary_goal:
                        lines.append(f"  - **Primary Goal:** {primary_goal}")

                    if success_metrics:
                        metrics_str = ", ".join(success_metrics[:3])
                        lines.append(f"  - **KPIs:** {metrics_str}")

                    metrics_line = []
                    if effort:
                        metrics_line.append(f"Effort: {effort}")
                    if roi:
                        metrics_line.append(f"ROI: {roi}")

                    if metrics_line:
                        lines.append(f"  - {' | '.join(metrics_line)}")

                lines.append("")

    # 5. Content Distribution Strategy
    content_dist = data.get("content_distribution", {})
    if content_dist and isinstance(content_dist, dict):
        lines.append("### Content Distribution & Repurposing")
        lines.append("")

        source = content_dist.get("source_platform", "")
        if source:
            source_name = str(source).replace("_", " ").title()
            lines.append(f"**Content Hub:** {source_name}")
            lines.append("Create your foundational content here, then adapt for other platforms.")
            lines.append("")

        flow = content_dist.get("distribution_flow", [])
        if flow:
            lines.append("**Distribution Flow:**")
            for step in flow[:5]:
                lines.append(f"  → {step}")
            lines.append("")

        repurposing = content_dist.get("repurposing_strategy", "")
        if repurposing:
            lines.append(f"**Repurposing Strategy:** {repurposing}")
            lines.append("")

        time_savings = content_dist.get("time_savings", "")
        if time_savings:
            lines.append(f"**Efficiency Gains:** {time_savings}")
            lines.append("")

    # 6. Current State Analysis (if provided)
    current_platforms = data.get("current_platforms", [])
    current_strengths = data.get("current_strengths", [])
    current_gaps = data.get("current_gaps", [])

    if current_platforms or current_strengths or current_gaps:
        lines.append("### Current State Analysis")
        lines.append("")

        if current_platforms:
            platforms_str = ", ".join(current_platforms)
            lines.append(f"**Currently Using:** {platforms_str}")
            lines.append("")

        if current_strengths:
            lines.append("**What's Working:**")
            for strength in current_strengths[:5]:
                lines.append(f"  ✓ {strength}")
            lines.append("")

        if current_gaps:
            lines.append("**What's Missing:**")
            for gap in current_gaps[:5]:
                lines.append(f"  ✗ {gap}")
            lines.append("")

    # 7. Quick Wins
    quick_wins = data.get("quick_wins", [])
    if quick_wins:
        lines.append("### Quick Wins - Start Here")
        lines.append("")

        for i, win in enumerate(quick_wins[:5], 1):
            if isinstance(win, dict):
                platform = str(win.get("platform", "Unknown")).replace("_", " ").title()
                action = win.get("action", "")
                timeframe = win.get("timeframe", "")
                outcome = win.get("expected_outcome", "")

                timeframe_label = f" ({timeframe})" if timeframe else ""
                lines.append(f"**{i}. {platform}**{timeframe_label}")

                if action:
                    lines.append(f"   **Action:** {action}")

                if outcome:
                    lines.append(f"   **Expected Outcome:** {outcome}")

                lines.append("")

    # 8. Implementation Plans
    thirty_day = data.get("thirty_day_plan", [])
    if thirty_day:
        lines.append("### 30-Day Implementation Plan")
        lines.append("")
        for i, action in enumerate(thirty_day[:7], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    ninety_day = data.get("ninety_day_plan", [])
    if ninety_day:
        lines.append("### 90-Day Strategic Rollout")
        lines.append("")
        for i, action in enumerate(ninety_day[:10], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 9. Strategic Insights
    key_insights = data.get("key_insights", [])
    if key_insights:
        lines.append("### Key Strategic Insights")
        lines.append("")
        for insight in key_insights[:5]:
            lines.append(f"- {insight}")
        lines.append("")

    # 10. Common Mistakes to Avoid
    mistakes = data.get("common_mistakes_to_avoid", [])
    if mistakes:
        lines.append("### Common Mistakes to Avoid")
        lines.append("")
        for mistake in mistakes[:5]:
            lines.append(f"⚠ {mistake}")
        lines.append("")

    return lines


def _format_audience_research(data: dict) -> List[str]:
    """Format audience research results with comprehensive audience insights.

    Displays demographics, psychographics, behavioral patterns, pain points, goals,
    content preferences, audience segments, messaging framework, and engagement tactics.
    """
    lines = []

    if not data:
        lines.append("**Audience Research:** No data available")
        return lines

    # 1. Executive Summary
    lines.append("### Audience Research Executive Summary")
    lines.append("")

    executive_summary = data.get("executive_summary", "")
    if executive_summary:
        lines.append(executive_summary)
        lines.append("")

    audience_size = data.get("audience_size_estimate", "")
    if audience_size:
        lines.append(f"**Total Addressable Audience:** {audience_size}")
        lines.append("")

    # 2. Demographics
    demographics = data.get("demographics", {})
    if demographics and isinstance(demographics, dict):
        lines.append("### Demographics")
        lines.append("")

        age_ranges = demographics.get("primary_age_ranges", [])
        if age_ranges:
            ages_str = ", ".join([str(age) for age in age_ranges])
            lines.append(f"**Age Ranges:** {ages_str}")

        gender = demographics.get("gender_distribution", "")
        if gender:
            lines.append(f"**Gender:** {gender}")

        locations = demographics.get("locations", [])
        if locations:
            locations_str = ", ".join(locations[:10])
            lines.append(f"**Locations:** {locations_str}")

        income_levels = demographics.get("income_levels", [])
        if income_levels:
            income_str = ", ".join([str(i).replace("_", " ").title() for i in income_levels])
            lines.append(f"**Income Levels:** {income_str}")

        education_levels = demographics.get("education_levels", [])
        if education_levels:
            edu_str = ", ".join([str(e).replace("_", " ").title() for e in education_levels])
            lines.append(f"**Education:** {edu_str}")

        job_titles = demographics.get("job_titles", [])
        if job_titles:
            lines.append("")
            lines.append("**Common Job Titles:**")
            for title in job_titles[:10]:
                lines.append(f"  - {title}")

        company_sizes = demographics.get("company_sizes", [])
        if company_sizes:
            sizes_str = ", ".join(company_sizes)
            lines.append("")
            lines.append(f"**Company Sizes:** {sizes_str}")

        lines.append("")

    # 3. Psychographics
    psychographics = data.get("psychographics", {})
    if psychographics and isinstance(psychographics, dict):
        lines.append("### Psychographics")
        lines.append("")

        values = psychographics.get("values", [])
        if values:
            lines.append("**Core Values:**")
            for value in values[:7]:
                lines.append(f"  - {value}")
            lines.append("")

        motivations = psychographics.get("motivations", [])
        if motivations:
            lines.append("**Key Motivations:**")
            for motivation in motivations[:7]:
                lines.append(f"  - {motivation}")
            lines.append("")

        personality_traits = psychographics.get("personality_traits", [])
        if personality_traits:
            traits_str = ", ".join(personality_traits[:7])
            lines.append(f"**Personality Traits:** {traits_str}")
            lines.append("")

        interests = psychographics.get("interests", [])
        if interests:
            interests_str = ", ".join(interests[:10])
            lines.append(f"**Interests:** {interests_str}")
            lines.append("")

        lifestyle = psychographics.get("lifestyle", "")
        if lifestyle:
            lines.append(f"**Lifestyle:** {lifestyle}")
            lines.append("")

    # 4. Behavioral Profile
    behavioral = data.get("behavioral_profile", {})
    if behavioral and isinstance(behavioral, dict):
        lines.append("### Behavioral Patterns")
        lines.append("")

        content_consumption = behavioral.get("content_consumption", "")
        if content_consumption:
            lines.append(f"**Content Consumption:** {content_consumption}")
            lines.append("")

        preferred_platforms = behavioral.get("preferred_platforms", [])
        if preferred_platforms:
            platforms_str = ", ".join(preferred_platforms[:10])
            lines.append(f"**Preferred Platforms:** {platforms_str}")
            lines.append("")

        online_behavior = behavioral.get("online_behavior", "")
        if online_behavior:
            lines.append(f"**Online Behavior:** {online_behavior}")
            lines.append("")

        purchase_behavior = behavioral.get("purchase_behavior", "")
        if purchase_behavior:
            lines.append(f"**Purchase Behavior:** {purchase_behavior}")
            lines.append("")

        engagement_patterns = behavioral.get("engagement_patterns", "")
        if engagement_patterns:
            lines.append(f"**Engagement Patterns:** {engagement_patterns}")
            lines.append("")

    # 5. Pain Points & Goals
    pain_points = data.get("pain_points", [])
    goals = data.get("goals_aspirations", [])

    if pain_points or goals:
        lines.append("### Pain Points & Goals")
        lines.append("")

        if pain_points:
            lines.append("**Top Pain Points:**")
            for i, pain in enumerate(pain_points[:10], 1):
                lines.append(f"{i}. {pain}")
            lines.append("")

        if goals:
            lines.append("**Goals & Aspirations:**")
            for i, goal in enumerate(goals[:10], 1):
                lines.append(f"{i}. {goal}")
            lines.append("")

    # 6. Content Preferences
    content_prefs = data.get("content_preferences", [])
    if content_prefs:
        lines.append("### Content Preferences")
        lines.append("")
        for pref in content_prefs[:10]:
            lines.append(f"- {pref}")
        lines.append("")

    # 7. Decision Factors
    decision_factors = data.get("decision_factors", [])
    if decision_factors:
        lines.append("### Decision-Making Factors")
        lines.append("")
        lines.append("What influences their buying decisions:")
        lines.append("")
        for i, factor in enumerate(decision_factors[:7], 1):
            lines.append(f"{i}. {factor}")
        lines.append("")

    # 8. Information Sources & Influencers
    info_sources = data.get("information_sources", [])
    influencers = data.get("influencers_brands", [])

    if info_sources or influencers:
        lines.append("### Information Sources & Influencers")
        lines.append("")

        if info_sources:
            lines.append("**Where They Get Information:**")
            for source in info_sources[:10]:
                lines.append(f"  - {source}")
            lines.append("")

        if influencers:
            lines.append("**Influencers & Brands They Follow:**")
            for influencer in influencers[:10]:
                lines.append(f"  - {influencer}")
            lines.append("")

    # 9. Audience Segments
    segments = data.get("audience_segments", [])
    if segments:
        lines.append("### Audience Segments")
        lines.append("")

        for i, segment in enumerate(segments[:5], 1):
            if isinstance(segment, dict):
                segment_name = segment.get("segment_name", "Unknown")
                segment_size = segment.get("segment_size", "")
                description = segment.get("description", "")
                characteristics = segment.get("key_characteristics", [])
                content_prefs = segment.get("content_preferences", [])
                messaging = segment.get("messaging_recommendations", "")

                size_label = f" ({segment_size})" if segment_size else ""
                lines.append(f"**{i}. {segment_name}**{size_label}")
                lines.append("")

                if description:
                    lines.append(description)
                    lines.append("")

                if characteristics:
                    lines.append("**Key Characteristics:**")
                    for char in characteristics[:5]:
                        lines.append(f"  - {char}")
                    lines.append("")

                if content_prefs:
                    lines.append("**Content Preferences:**")
                    for pref in content_prefs[:5]:
                        lines.append(f"  - {pref}")
                    lines.append("")

                if messaging:
                    lines.append(f"**Messaging:** {messaging}")
                    lines.append("")

    # 10. Messaging Framework
    messaging_framework = data.get("messaging_framework", "")
    if messaging_framework:
        lines.append("### Messaging Framework")
        lines.append("")
        lines.append(messaging_framework)
        lines.append("")

    # 11. Content Strategy Recommendations
    content_strategy = data.get("content_strategy_recommendations", [])
    if content_strategy:
        lines.append("### Content Strategy Recommendations")
        lines.append("")
        for i, rec in enumerate(content_strategy[:7], 1):
            lines.append(f"{i}. {rec}")
        lines.append("")

    # 12. Engagement Tactics
    engagement_tactics = data.get("engagement_tactics", [])
    if engagement_tactics:
        lines.append("### Engagement Tactics")
        lines.append("")
        lines.append("How to engage this audience effectively:")
        lines.append("")
        for i, tactic in enumerate(engagement_tactics[:7], 1):
            lines.append(f"{i}. {tactic}")
        lines.append("")

    # 13. Key Insights
    key_insights = data.get("key_insights", [])
    if key_insights:
        lines.append("### Key Strategic Insights")
        lines.append("")
        for insight in key_insights[:7]:
            lines.append(f"- {insight}")
        lines.append("")

    # 14. What to Avoid
    what_to_avoid = data.get("what_to_avoid", [])
    if what_to_avoid:
        lines.append("### What to Avoid")
        lines.append("")
        for avoid in what_to_avoid[:5]:
            lines.append(f"⚠ {avoid}")
        lines.append("")

    return lines


def _format_content_calendar(data: dict) -> List[str]:
    """Format content calendar results with comprehensive 90-day plan.

    Displays executive summary, content pillars, themes, weekly calendar, platform schedules,
    posting strategy, seasonal opportunities, implementation guidance, and success metrics.
    """
    lines = []

    if not data:
        lines.append("**Content Calendar:** No data available")
        return lines

    # 1. Executive Summary
    lines.append("### Content Calendar Strategy")
    lines.append("")

    start_date = data.get("strategy_start_date", "")
    if start_date:
        lines.append(f"**Calendar Start:** {start_date}")

    total_posts = data.get("total_posts_90_days", 0)
    if total_posts:
        lines.append(f"**Total Posts (90 Days):** {total_posts}")

    lines.append("")

    executive_summary = data.get("executive_summary", "")
    if executive_summary:
        lines.append(executive_summary)
        lines.append("")

    # Primary goals
    primary_goals = data.get("primary_goals", [])
    if primary_goals:
        goals_str = ", ".join([str(g).replace("_", " ").title() for g in primary_goals])
        lines.append(f"**Primary Goals:** {goals_str}")
        lines.append("")

    # 2. Content Pillars
    content_pillars = data.get("content_pillars", [])
    pillar_rationale = data.get("pillar_rationale", "")

    if content_pillars or pillar_rationale:
        lines.append("### Content Pillars")
        lines.append("")

        if content_pillars:
            pillars_str = ", ".join([str(p).replace("_", " ").title() for p in content_pillars])
            lines.append(f"**Pillars:** {pillars_str}")
            lines.append("")

        if pillar_rationale:
            lines.append(f"**Strategy:** {pillar_rationale}")
            lines.append("")

    # 3. Content Mix
    content_mix = data.get("content_mix", "")
    if content_mix:
        lines.append("### Content Mix")
        lines.append("")
        lines.append(content_mix)
        lines.append("")

    # 4. Themes
    themes = data.get("themes", [])
    if themes:
        lines.append("### Content Themes")
        lines.append("")

        for i, theme in enumerate(themes[:3], 1):
            if isinstance(theme, dict):
                name = theme.get("name", "Unknown")
                description = theme.get("description", "")
                goal = theme.get("goal", "")
                weeks = theme.get("weeks", [])

                lines.append(f"**{i}. {name}**")

                if description:
                    lines.append(f"   {description}")

                if goal:
                    goal_name = str(goal).replace("_", " ").title()
                    lines.append(f"   - **Goal:** {goal_name}")

                if weeks:
                    weeks_str = ", ".join([str(w) for w in weeks])
                    lines.append(f"   - **Weeks:** {weeks_str}")

                lines.append("")

    # 5. Weekly Calendar (First 4 weeks detail, then summary)
    weekly_calendar = data.get("weekly_calendar", [])
    if weekly_calendar:
        lines.append("### 90-Day Weekly Calendar")
        lines.append("")

        # Show first 4 weeks in detail
        for week in weekly_calendar[:4]:
            if isinstance(week, dict):
                week_num = week.get("week_number", 0)
                start = week.get("start_date", "")
                theme = week.get("theme", "")
                pillar = str(week.get("pillar", "")).replace("_", " ").title()
                post_count = week.get("post_count", 0)
                goal = str(week.get("goal", "")).replace("_", " ").title()
                key_message = week.get("key_message", "")
                post_ideas = week.get("post_ideas", [])
                holidays = week.get("holidays_events", [])

                lines.append(f"**Week {week_num}** ({start}) - {theme}")
                lines.append(
                    f"  - **Pillar:** {pillar} | **Posts:** {post_count} | **Goal:** {goal}"
                )

                if key_message:
                    lines.append(f"  - **Message:** {key_message}")

                if post_ideas:
                    lines.append("  - **Post Ideas:**")
                    for idea in post_ideas[:3]:
                        lines.append(f"    - {idea}")

                if holidays:
                    holidays_str = ", ".join(holidays)
                    lines.append(f"  - **Events:** {holidays_str}")

                lines.append("")

        # Summarize remaining weeks
        if len(weekly_calendar) > 4:
            lines.append(f"*...and {len(weekly_calendar) - 4} more weeks planned*")
            lines.append("")

    # 6. Platform Calendars
    platform_calendars = data.get("platform_calendars", [])
    if platform_calendars:
        lines.append("### Platform Posting Schedule")
        lines.append("")

        for platform_cal in platform_calendars[:5]:
            if isinstance(platform_cal, dict):
                platform = platform_cal.get("platform", "Unknown")
                frequency = str(platform_cal.get("frequency", "")).replace("_", " ").title()
                best_days = platform_cal.get("best_days", [])
                best_times = platform_cal.get("best_times", [])
                content_mix = platform_cal.get("content_mix", "")

                lines.append(f"**{platform}** ({frequency})")

                if best_days:
                    days_str = ", ".join(best_days[:3])
                    lines.append(f"  - **Best Days:** {days_str}")

                if best_times:
                    times_str = ", ".join(best_times[:3])
                    lines.append(f"  - **Best Times:** {times_str}")

                if content_mix:
                    lines.append(f"  - **Mix:** {content_mix}")

                lines.append("")

    # 7. Seasonal Opportunities
    seasonal_opps = data.get("seasonal_opportunities", [])
    if seasonal_opps:
        lines.append("### Seasonal Content Opportunities")
        lines.append("")
        for i, opp in enumerate(seasonal_opps[:10], 1):
            lines.append(f"{i}. {opp}")
        lines.append("")

    # 8. Quick Start Actions
    quick_start = data.get("quick_start_actions", [])
    if quick_start:
        lines.append("### Quick Start (Week 1)")
        lines.append("")
        for i, action in enumerate(quick_start[:5], 1):
            lines.append(f"{i}. {action}")
        lines.append("")

    # 9. Content Creation Workflow
    workflow = data.get("content_creation_workflow", "")
    if workflow:
        lines.append("### Content Creation Workflow")
        lines.append("")
        lines.append(workflow)
        lines.append("")

    # 10. Batch Creation Tips
    batch_tips = data.get("batch_creation_tips", [])
    if batch_tips:
        lines.append("### Batch Creation Tips")
        lines.append("")
        for tip in batch_tips[:5]:
            lines.append(f"- {tip}")
        lines.append("")

    # 11. Success Metrics
    success_metrics = data.get("success_metrics", [])
    if success_metrics:
        lines.append("### Success Metrics")
        lines.append("")
        lines.append("Track these KPIs:")
        lines.append("")
        for metric in success_metrics[:7]:
            lines.append(f"- {metric}")
        lines.append("")

    # 12. Review Schedule
    review_schedule = data.get("review_schedule", "")
    if review_schedule:
        lines.append(f"**Review Schedule:** {review_schedule}")
        lines.append("")

    # 13. Key Insights
    key_insights = data.get("key_insights", [])
    if key_insights:
        lines.append("### Key Strategic Insights")
        lines.append("")
        for insight in key_insights[:5]:
            lines.append(f"- {insight}")
        lines.append("")

    # 14. Common Pitfalls
    common_pitfalls = data.get("common_pitfalls", [])
    if common_pitfalls:
        lines.append("### Common Pitfalls to Avoid")
        lines.append("")
        for pitfall in common_pitfalls[:5]:
            lines.append(f"⚠ {pitfall}")
        lines.append("")

    return lines


def _format_icp_workshop(data: dict) -> List[str]:
    """Format ICP workshop results with comprehensive ideal customer profile.

    Displays demographics, psychographics, behavioral patterns, situational factors,
    success criteria, and actionable ICP summary.
    """
    lines = []

    if not data:
        lines.append("**ICP Workshop:** No data available")
        return lines

    lines.append("### Ideal Customer Profile")
    lines.append("")

    # Profile name
    profile_name = data.get("profile_name", "")
    if profile_name:
        lines.append(f"**Profile:** {profile_name}")
        lines.append("")

    # One-sentence summary
    summary = data.get("one_sentence_summary", "")
    if summary:
        lines.append(f"**Summary:** {summary}")
        lines.append("")

    # Demographics
    demographics = data.get("demographics", {})
    if demographics:
        lines.append("### Demographics/Firmographics")
        lines.append("")

        for key in ["company_size", "industry", "revenue_range", "location", "team_structure"]:
            value = demographics.get(key)
            if value:
                label = key.replace("_", " ").title()
                lines.append(f"**{label}:** {value}")

        job_titles = demographics.get("job_titles", [])
        if job_titles:
            lines.append(f"**Decision Makers:** {', '.join(job_titles[:5])}")

        tech = demographics.get("technologies_used", [])
        if tech:
            lines.append(f"**Tech Stack:** {', '.join(tech[:5])}")

        lines.append("")

    # Psychographics
    psycho = data.get("psychographics", {})
    if psycho:
        lines.append("### Psychographics")
        lines.append("")

        goals = psycho.get("goals", [])
        if goals:
            lines.append("**Goals:**")
            for goal in goals[:5]:
                lines.append(f"  - {goal}")
            lines.append("")

        challenges = psycho.get("challenges", [])
        if challenges:
            lines.append("**Challenges:**")
            for challenge in challenges[:5]:
                lines.append(f"  - {challenge}")
            lines.append("")

        values = psycho.get("values", [])
        if values:
            lines.append(f"**Values:** {', '.join(values[:5])}")
            lines.append("")

    # Behavioral
    behavioral = data.get("behavioral", {})
    if behavioral:
        lines.append("### Behavioral Patterns")
        lines.append("")

        for key in ["buying_process", "research_habits"]:
            value = behavioral.get(key)
            if value:
                label = key.replace("_", " ").title()
                lines.append(f"**{label}:** {value}")

        platforms = behavioral.get("platforms_active_on", [])
        if platforms:
            lines.append(f"**Active Platforms:** {', '.join(platforms[:5])}")

        lines.append("")

    # Success criteria
    success = data.get("success_criteria", {})
    if success:
        lines.append("### Success Criteria")
        lines.append("")

        definition = success.get("definition_of_success")
        if definition:
            lines.append(f"**Definition:** {definition}")

        kpis = success.get("kpis_tracked", [])
        if kpis:
            lines.append(f"**KPIs:** {', '.join(kpis[:5])}")

        lines.append("")

    return lines


def _format_story_mining(data: dict) -> List[str]:
    """Format story mining results with customer success stories."""
    lines = []

    if not data:
        lines.append("**Story Mining:** No data available")
        return lines

    lines.append("### Customer Success Stories")
    lines.append("")

    # Story summary
    story_count = data.get("total_stories", 0)
    if story_count:
        lines.append(f"**Stories Identified:** {story_count}")
        lines.append("")

    # Customer journey
    journey = data.get("customer_journey", {})
    if journey:
        lines.append("### Customer Journey")
        lines.append("")

        before = journey.get("before", {})
        if before:
            lines.append("**Before (Pain State):**")
            situation = before.get("situation", "")
            if situation:
                lines.append(f"  - Situation: {situation}")
            pain_points = before.get("pain_points", [])
            for pain in pain_points[:3]:
                lines.append(f"  - {pain}")
            lines.append("")

        decision = journey.get("decision", {})
        if decision:
            lines.append("**Decision Point:**")
            trigger = decision.get("trigger", "")
            if trigger:
                lines.append(f"  - Trigger: {trigger}")
            lines.append("")

        after = journey.get("after", {})
        if after:
            lines.append("**After (Success State):**")
            results = after.get("results", [])
            for result in results[:5]:
                lines.append(f"  - {result}")
            lines.append("")

    # Key quotes
    quotes = data.get("key_quotes", [])
    if quotes:
        lines.append("### Key Customer Quotes")
        lines.append("")
        for quote in quotes[:5]:
            lines.append(f'> "{quote}"')
            lines.append("")

    # Story angles
    angles = data.get("story_angles", [])
    if angles:
        lines.append("### Story Angles")
        lines.append("")
        for angle in angles[:5]:
            lines.append(f"- {angle}")
        lines.append("")

    return lines


def _format_market_trends(data: dict) -> List[str]:
    """Format market trends results with comprehensive trend analysis.

    Displays market summary, rising/relevant trends with momentum and growth rates,
    emerging conversations, seasonal patterns, content opportunities, and key themes.
    """
    lines = []

    # 1. Market Summary (executive overview)
    if data.get("market_summary"):
        lines.append("### Market Trends Summary")
        lines.append("")
        lines.append(data["market_summary"])
        lines.append("")

    # 2. Top Rising Trends (fastest growing)
    top_rising = data.get("top_rising_trends", [])
    if top_rising:
        lines.append("### Top Rising Trends")
        lines.append("")
        lines.append("| # | Trend | Momentum | Relevance | Growth | Popularity | Urgency |")
        lines.append("|---|-------|----------|-----------|--------|------------|---------|")

        for i, trend in enumerate(top_rising[:5], 1):
            topic = trend.get("topic", "Unknown")
            momentum = trend.get("momentum", "unknown").upper()
            relevance = trend.get("relevance", "unknown").upper()
            growth = trend.get("growth_rate", "N/A")
            popularity = trend.get("popularity_score", 0)
            urgency = trend.get("urgency", "Medium")

            # Add emoji for momentum
            momentum_emoji = {
                "RISING": "↗",
                "EMERGING": "✦",
                "STABLE": "→",
                "DECLINING": "↘",
                "SEASONAL": "~",
            }.get(momentum, "")

            lines.append(
                f"| {i} | {topic} | {momentum} {momentum_emoji} | {relevance} | "
                f"{growth} | {popularity:.1f}/10 | {urgency} |"
            )

        lines.append("")

        # Show detailed info for top 3 trends
        lines.append("**Trend Details:**")
        lines.append("")
        for i, trend in enumerate(top_rising[:3], 1):
            lines.append(f"**{i}. {trend.get('topic', 'Unknown')}**")

            description = trend.get("description", "")
            if description:
                lines.append(f"_{description}_")
                lines.append("")

            key_drivers = trend.get("key_drivers", [])
            if key_drivers:
                lines.append(f"- **Drivers:** {', '.join(key_drivers[:3])}")

            target_audience = trend.get("target_audience", [])
            if target_audience:
                lines.append(f"- **Audience:** {', '.join(target_audience[:3])}")

            content_angles = trend.get("content_angles", [])
            if content_angles:
                lines.append(f"- **Content angles:** {', '.join(content_angles[:2])}")

            related_keywords = trend.get("related_keywords", [])
            if related_keywords:
                lines.append(f"- **Keywords:** {', '.join(related_keywords[:5])}")

            lines.append("")

    # 3. Immediate Opportunities (capitalize now)
    immediate_opps = data.get("immediate_opportunities", [])
    if immediate_opps:
        lines.append("### Immediate Content Opportunities")
        lines.append("")
        lines.append("Capitalize on these trends NOW:")
        lines.append("")
        for i, opp in enumerate(immediate_opps, 1):
            lines.append(f"{i}. {opp}")
        lines.append("")

    # 4. Emerging Conversations (new industry debates)
    emerging = data.get("emerging_conversations", [])
    if emerging:
        lines.append("### Emerging Industry Conversations")
        lines.append("")

        for conv in emerging[:3]:
            topic = conv.get("topic", "Unknown")
            description = conv.get("description", "")

            lines.append(f"**{topic}**")
            if description:
                lines.append(f"_{description}_")

            perspectives = conv.get("key_perspectives", [])
            if perspectives:
                lines.append(f"- Key viewpoints: {', '.join(perspectives[:3])}")

            thought_leaders = conv.get("thought_leaders", [])
            if thought_leaders:
                lines.append(f"- Thought leaders: {', '.join(thought_leaders[:3])}")

            opportunity = conv.get("content_opportunity", "")
            if opportunity:
                lines.append(f"- **Opportunity:** {opportunity}")

            lines.append("")

    # 5. Seasonal Trends (recurring patterns)
    seasonal = data.get("seasonal_trends", [])
    if seasonal:
        lines.append("### Seasonal Trends")
        lines.append("")

        for trend in seasonal[:5]:
            topic = trend.get("topic", "Unknown")
            timing = trend.get("timing", "Unknown")
            prep = trend.get("preparation_timeline", "")

            lines.append(f"**{topic}** ({timing})")

            description = trend.get("description", "")
            if description:
                lines.append(f"- {description}")

            if prep:
                lines.append(f"- **Start prep:** {prep}")

            lines.append("")

    # 6. Upcoming Opportunities (prepare for)
    upcoming_opps = data.get("upcoming_opportunities", [])
    if upcoming_opps:
        lines.append("### Upcoming Opportunities")
        lines.append("")
        lines.append("Prepare content for these trends:")
        lines.append("")
        for i, opp in enumerate(upcoming_opps, 1):
            lines.append(f"{i}. {opp}")
        lines.append("")

    # 7. Key Themes (overarching narratives)
    key_themes = data.get("key_themes", [])
    if key_themes:
        lines.append("### Key Market Themes")
        lines.append("")
        for theme in key_themes:
            lines.append(f"- {theme}")
        lines.append("")

    # 8. Declining Topics (what to avoid)
    declining = data.get("declining_topics", [])
    if declining:
        lines.append("### Declining Topics")
        lines.append("")
        lines.append("Avoid creating content on these topics (losing relevance):")
        lines.append("")
        for topic in declining:
            lines.append(f"- {topic}")
        lines.append("")

    # If no data available
    if not any(
        [
            top_rising,
            immediate_opps,
            emerging,
            seasonal,
            upcoming_opps,
            key_themes,
            declining,
        ]
    ):
        lines.append("**Market Trends:** No data available")

    return lines


def _format_determine_competitors(data: dict) -> List[str]:
    """Format determine competitors results for export"""
    lines = []

    if "primary_competitors" in data and data["primary_competitors"]:
        lines.append("**Identified Competitors:**")
        lines.append("")
        for comp in data["primary_competitors"][:5]:
            name = comp.get("name", "Unknown")
            threat = comp.get("threat_level", "medium").upper()
            position = comp.get("market_position", "N/A")
            diff = comp.get("differentiation_opportunity", "")

            lines.append(f"### {name} ({threat} Threat)")
            lines.append(f"**Position:** {position}")
            if diff:
                lines.append(f"**How to Differentiate:** {diff}")
            lines.append("")

    # Positioning recommendation
    if "recommended_positioning" in data:
        lines.append("**Recommended Positioning:**")
        lines.append(data["recommended_positioning"])
        lines.append("")

    return lines


def _format_generic_research(data: dict) -> List[str]:
    """Generic fallback formatter for research data."""
    lines = []
    for key, value in data.items():
        if isinstance(value, list):
            lines.append(f"**{key.replace('_', ' ').title()}:**")
            for item in value[:10]:  # Limit to 10 items
                lines.append(f"- {item}")
            lines.append("")
        elif isinstance(value, dict):
            lines.append(f"**{key.replace('_', ' ').title()}:**")
            for subkey, subval in value.items():
                lines.append(f"- {subkey}: {subval}")
            lines.append("")
        else:
            lines.append(f"**{key.replace('_', ' ').title()}:** {value}")
    return lines
